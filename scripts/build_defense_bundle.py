"""
根据两份答辩 Markdown 资料,合并生成 Word 和 PDF。

源文件:
  - presentation/答辩讲稿_10min_基于PDF.md   (讲稿)
  - presentation/答辩Q&A_范式与调参.md        (Q&A 弹药库)

输出:
  - presentation/OpsWarden_答辩资料合集.docx
  - presentation/OpsWarden_答辩资料合集.pdf
"""

import re
from pathlib import Path
from typing import List, Tuple, Optional

# ============================================================
# 配置
# ============================================================
REPO = Path("d:/Repositories/OpsWarden")
SRC_SCRIPT = REPO / "presentation" / "答辩讲稿_10min_基于PDF.md"
SRC_QA = REPO / "presentation" / "答辩Q&A_范式与调参.md"
OUT_DOCX = REPO / "presentation" / "OpsWarden_答辩资料合集.docx"
OUT_PDF = REPO / "presentation" / "OpsWarden_答辩资料合集.pdf"
FONT_PATH = "C:/Windows/Fonts/simhei.ttf"  # 黑体(单一 TTF,无 subset 警告)


# ============================================================
# Markdown 解析:把 MD 转成中间结构
# ============================================================
# 块类型: heading, paragraph, quote, table, hr, code, blank
Block = dict  # {'type': 'heading'|'paragraph'|'quote'|'table'|'hr'|'code', ...}


def parse_markdown(md_text: str) -> List[Block]:
    """极简 Markdown 解析(够用即可)。"""
    lines = md_text.split("\n")
    blocks: List[Block] = []
    i = 0
    n = len(lines)
    in_code = False
    code_buf: List[str] = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块围栏
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                blocks.append({"type": "code", "content": "\n".join(code_buf)})
                in_code = False
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 水平线
        if re.match(r"^[-*_]{3,}\s*$", stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # 引用块(连续以 > 开头,聚合)
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                content = re.sub(r"^>\s?", "", lines[i].strip())
                quote_lines.append(content)
                i += 1
            blocks.append({"type": "quote", "content": "\n".join(quote_lines)})
            continue

        # 表格(GFM 风格: | col | col |)
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # 过滤分隔行(像 |---|---|)
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(re.match(r"^[-:\s]+$", c) for c in cells):
                    continue
                rows.append(cells)
            if rows:
                # 列数对齐
                max_cols = max(len(r) for r in rows)
                rows = [r + [""] * (max_cols - len(r)) for r in rows]
                blocks.append({"type": "table", "rows": rows})
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 段落(聚合连续非空非特殊行)
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if not nxt or nxt.startswith("#") or nxt.startswith(">") or nxt.startswith("|") or nxt.startswith("```") or re.match(r"^[-*_]{3,}\s*$", nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def parse_inline(text: str) -> List[Tuple[str, bool]]:
    """把行内 **粗体** 拆成 [(text, is_bold), ...]。"""
    parts: List[Tuple[str, bool]] = []
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False))
        parts.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False))
    if not parts:
        parts = [(text, False)]
    return parts


# ============================================================
# 生成 Word (python-docx)
# ============================================================
def gen_docx(all_blocks: List[Tuple[str, List[Block]]], out_path: Path):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn

    doc = Document()

    # 全局样式:宋体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    rpr = style.element.rPr
    if rpr is None:
        from docx.oxml import OxmlElement
        rpr = OxmlElement("w:rPr")
        style.element.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 封面
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("OpsWarden")
    title_run.font.name = "微软雅黑"
    title_run.font.size = Pt(36)
    title_run.bold = True
    title_rpr = title_run._element.get_or_add_rPr()
    title_rfonts = title_rpr.find(qn("w:rFonts"))
    if title_rfonts is None:
        from docx.oxml import OxmlElement
        title_rfonts = OxmlElement("w:rFonts")
        title_rpr.append(title_rfonts)
    title_rfonts.set(qn("w:eastAsia"), "微软雅黑")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("运维数字员工 · 答辩资料合集")
    sub_run.font.name = "微软雅黑"
    sub_run.font.size = Pt(18)
    sub_run.bold = False
    sub_rpr = sub_run._element.get_or_add_rPr()
    sub_rfonts = sub_rpr.find(qn("w:rFonts"))
    if sub_rfonts is None:
        from docx.oxml import OxmlElement
        sub_rfonts = OxmlElement("w:rFonts")
        sub_rpr.append(sub_rfonts)
    sub_rfonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("团队:廖晨扬 · 吴雨彤 · 丁其彬  |  2023 级计算机科学与技术 2 班")
    info_run.font.name = "宋体"
    info_run.font.size = Pt(12)
    info_rpr = info_run._element.get_or_add_rPr()
    info_rfonts = info_rpr.find(qn("w:rFonts"))
    if info_rfonts is None:
        from docx.oxml import OxmlElement
        info_rfonts = OxmlElement("w:rFonts")
        info_rpr.append(info_rfonts)
    info_rfonts.set(qn("w:eastAsia"), "宋体")

    doc.add_paragraph()
    tip_p = doc.add_paragraph()
    tip_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tip_run = tip_p.add_run("本文档包含两部分:\n  第一部分:10 分钟答辩讲稿(贴 PDF 演示)\n  第二部分:范式对比 + 超参调参 Q&A 弹药库(30 问 + 速记卡)")
    tip_run.font.name = "宋体"
    tip_run.font.size = Pt(11)
    tip_rpr = tip_run._element.get_or_add_rPr()
    tip_rfonts = tip_rpr.find(qn("w:rFonts"))
    if tip_rfonts is None:
        from docx.oxml import OxmlElement
        tip_rfonts = OxmlElement("w:rFonts")
        tip_rpr.append(tip_rfonts)
    tip_rfonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()

    for src_name, blocks in all_blocks:
        # 分节大标题
        h0 = doc.add_paragraph()
        h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h0_run = h0.add_run(src_name)
        h0_run.font.name = "微软雅黑"
        h0_run.font.size = Pt(20)
        h0_run.bold = True
        h0_rpr = h0_run._element.get_or_add_rPr()
        h0_rfonts = h0_rpr.find(qn("w:rFonts"))
        if h0_rfonts is None:
            from docx.oxml import OxmlElement
            h0_rfonts = OxmlElement("w:rFonts")
            h0_rpr.append(h0_rfonts)
        h0_rfonts.set(qn("w:eastAsia"), "微软雅黑")
        # 蓝下划线
        from docx.oxml import OxmlElement
        pPr = h0._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.add_paragraph()

        for blk in blocks:
            t = blk["type"]
            if t == "heading":
                level = blk["level"]
                p = doc.add_paragraph()
                run = p.add_run(blk["text"])
                if level == 1:
                    run.font.size = Pt(18)
                    run.bold = True
                    run.font.name = "微软雅黑"
                    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                elif level == 2:
                    run.font.size = Pt(15)
                    run.bold = True
                    run.font.name = "微软雅黑"
                    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                elif level == 3:
                    run.font.size = Pt(13)
                    run.bold = True
                    run.font.name = "微软雅黑"
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                else:
                    run.font.size = Pt(12)
                    run.bold = True
                    run.font.name = "微软雅黑"
                run_rpr = run._element.get_or_add_rPr()
                run_rfonts = run_rpr.find(qn("w:rFonts"))
                if run_rfonts is None:
                    from docx.oxml import OxmlElement
                    run_rfonts = OxmlElement("w:rFonts")
                    run_rpr.append(run_rfonts)
                run_rfonts.set(qn("w:eastAsia"), "微软雅黑")

            elif t == "paragraph":
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.line_spacing = 1.5
                for seg_text, is_bold in parse_inline(blk["text"]):
                    run = p.add_run(seg_text)
                    run.font.size = Pt(11)
                    run.bold = is_bold
                    run.font.name = "宋体"
                    run_rpr = run._element.get_or_add_rPr()
                    run_rfonts = run_rpr.find(qn("w:rFonts"))
                    if run_rfonts is None:
                        from docx.oxml import OxmlElement
                        run_rfonts = OxmlElement("w:rFonts")
                        run_rpr.append(run_rfonts)
                    run_rfonts.set(qn("w:eastAsia"), "宋体")

            elif t == "quote":
                # 引用块:多行合一或分行
                for ql in blk["content"].split("\n"):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.74)
                    p.paragraph_format.right_indent = Cm(0.74)
                    p.paragraph_format.line_spacing = 1.4
                    # 左侧竖线
                    from docx.oxml import OxmlElement
                    pPr = p._element.get_or_add_pPr()
                    pBdr = OxmlElement("w:pBdr")
                    left = OxmlElement("w:left")
                    left.set(qn("w:val"), "single")
                    left.set(qn("w:sz"), "18")
                    left.set(qn("w:space"), "4")
                    left.set(qn("w:color"), "2E75B6")
                    pBdr.append(left)
                    pPr.append(pBdr)
                    # 浅灰底
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F2F8FC")
                    pPr.append(shd)
                    for seg_text, is_bold in parse_inline(ql):
                        run = p.add_run(seg_text)
                        run.font.size = Pt(11)
                        run.bold = is_bold
                        run.italic = False
                        run.font.name = "宋体"
                        run_rpr = run._element.get_or_add_rPr()
                        run_rfonts = run_rpr.find(qn("w:rFonts"))
                        if run_rfonts is None:
                            run_rfonts = OxmlElement("w:rFonts")
                            run_rpr.append(run_rfonts)
                        run_rfonts.set(qn("w:eastAsia"), "宋体")

            elif t == "table":
                rows = blk["rows"]
                if not rows:
                    continue
                n_cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Light Grid Accent 1"
                # 紧凑一点
                table.autofit = True
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        cell = table.rows[ri].cells[ci]
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        # 清空默认段落
                        cell.text = ""
                        p = cell.paragraphs[0]
                        for seg_text, is_bold in parse_inline(cell_text):
                            run = p.add_run(seg_text)
                            run.font.size = Pt(10)
                            run.bold = is_bold or (ri == 0)
                            run.font.name = "宋体"
                            run_rpr = run._element.get_or_add_rPr()
                            run_rfonts = run_rpr.find(qn("w:rFonts"))
                            if run_rfonts is None:
                                from docx.oxml import OxmlElement
                                run_rfonts = OxmlElement("w:rFonts")
                                run_rpr.append(run_rfonts)
                            run_rfonts.set(qn("w:eastAsia"), "宋体")

            elif t == "hr":
                p = doc.add_paragraph()
                from docx.oxml import OxmlElement
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "6")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "AAAAAA")
                pBdr.append(bottom)
                pPr.append(pBdr)

            elif t == "code":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(blk["content"])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # 每个分节后插分页
        doc.add_page_break()

    doc.save(str(out_path))
    print(f"[OK] Word 已生成: {out_path}")


# ============================================================
# 生成 PDF (fpdf2)
# ============================================================
def gen_pdf(all_blocks: List[Tuple[str, List[Block]]], out_path: Path, font_path: str):
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)
    # 注册中文字体
    pdf.add_font("zh", "", font_path)
    pdf.add_font("zh", "B", font_path)

    # 可用宽度常量(fpdf2 multi_cell 不接受 width=0)
    USABLE = lambda: pdf.w - pdf.l_margin - pdf.r_margin
    L_M = lambda: pdf.l_margin
    R_M = lambda: pdf.r_margin

    # 封面
    pdf.add_page()
    pdf.set_font("zh", "B", 36)
    pdf.ln(40)
    pdf.cell(0, 18, "OpsWarden", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("zh", "", 16)
    pdf.cell(0, 12, "运维数字员工 · 答辩资料合集", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(20)
    pdf.set_font("zh", "", 12)
    pdf.cell(0, 8, "团队:廖晨扬 · 吴雨彤 · 丁其彬", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(0, 8, "2023 级计算机科学与技术 2 班", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(20)
    pdf.set_font("zh", "", 11)
    pdf.cell(0, 7, "本文档包含两部分:", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(0, 7, "  第一部分:10 分钟答辩讲稿(贴 PDF 演示)", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(0, 7, "  第二部分:范式对比 + 超参调参 Q&A 弹药库(30 问 + 速记卡)", new_x="LMARGIN", new_y="NEXT", align="C")

    # 正文
    for src_name, blocks in all_blocks:
        pdf.add_page()
        # 分节大标题
        pdf.set_font("zh", "B", 18)
        pdf.set_text_color(46, 117, 182)
        pdf.ln(8)
        pdf.cell(0, 12, src_name, new_x="LMARGIN", new_y="NEXT", align="C")
        # 标题下划线
        y = pdf.get_y()
        pdf.set_draw_color(46, 117, 182)
        pdf.set_line_width(0.6)
        pdf.line(20, y, 190, y)
        pdf.ln(6)
        pdf.set_text_color(0, 0, 0)

        for blk in blocks:
            t = blk["type"]
            if t == "heading":
                level = blk["level"]
                if level == 1:
                    pdf.set_font("zh", "B", 16)
                    pdf.set_text_color(46, 117, 182)
                    pdf.ln(4)
                elif level == 2:
                    pdf.set_font("zh", "B", 14)
                    pdf.set_text_color(46, 117, 182)
                    pdf.ln(3)
                elif level == 3:
                    pdf.set_font("zh", "B", 12)
                    pdf.set_text_color(31, 78, 121)
                else:
                    pdf.set_font("zh", "B", 11)
                    pdf.set_text_color(31, 78, 121)
                # 多行支持
                pdf.multi_cell(USABLE(), 8, _strip_md(blk["text"]))
                pdf.set_text_color(0, 0, 0)

            elif t == "paragraph":
                pdf.set_font("zh", "", 10.5)
                # 解析粗体: 简化为只显示文本
                pdf.multi_cell(USABLE(), 6, _strip_md(blk["text"]))

            elif t == "quote":
                pdf.set_font("zh", "", 10.5)
                pdf.set_text_color(60, 60, 60)
                inner = USABLE() - 12
                for ql in blk["content"].split("\n"):
                    pdf.set_x(L_M() + 6)
                    pdf.multi_cell(inner, 6, _strip_md(ql))
                pdf.set_text_color(0, 0, 0)

            elif t == "table":
                rows = blk["rows"]
                if not rows:
                    continue
                n_cols = len(rows[0])
                col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / n_cols
                # 头行底色
                pdf.set_font("zh", "B", 9)
                pdf.set_fill_color(213, 232, 240)
                for ci, cell_text in enumerate(rows[0]):
                    pdf.cell(col_w, 7, _strip_md(cell_text), border=1, fill=True, align="C")
                pdf.ln()
                pdf.set_font("zh", "", 9)
                for ri, row in enumerate(rows[1:], start=1):
                    if ri % 2 == 0:
                        pdf.set_fill_color(245, 248, 250)
                    else:
                        pdf.set_fill_color(255, 255, 255)
                    for ci, cell_text in enumerate(row):
                        # 多行单元格:auto
                        lines = _strip_md(cell_text).split("\n")
                        n_lines = max(1, len(lines))
                        h = 6 * n_lines + 2
                        # 用 multi_cell 不支持同表格列对齐, 简化为单行截断
                        txt = _strip_md(cell_text)[:60]
                        pdf.cell(col_w, h, txt, border=1, fill=(ri % 2 == 0))
                    pdf.ln()

            elif t == "hr":
                y = pdf.get_y() + 2
                pdf.set_draw_color(180, 180, 180)
                pdf.set_line_width(0.3)
                pdf.line(20, y, 190, y)
                pdf.ln(4)

            elif t == "code":
                pdf.set_font("Courier", "", 9)
                pdf.set_text_color(50, 50, 50)
                pdf.multi_cell(USABLE(), 5, blk["content"])
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("zh", "", 10.5)

    pdf.output(str(out_path))
    print(f"[OK] PDF 已生成: {out_path}")


def _strip_md(text: str) -> str:
    """把简单的 markdown 行内标记去掉,并把 emoji / 半角 ¥ 替换为 PDF 友好字符。"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("❌", "[X]").replace("✅", "[V]").replace("✓", "[v]").replace("★", "*")
    text = text.replace("¥", "￥")  # 全角人民币符号,SimHei 支持
    return text


def _replace_emoji(text: str) -> str:
    """把 emoji 替换为 ASCII 标记(Word 端用,保留 ** 供粗体解析)。"""
    return (text.replace("❌", "[X]").replace("✅", "[V]")
                .replace("✓", "[v]").replace("★", "*"))


def _clean_blocks(blocks: List[Block]) -> List[Block]:
    """递归把所有 block 文本里的 emoji 替换为 ASCII。"""
    for b in blocks:
        t = b["type"]
        if t in ("heading", "paragraph"):
            b["text"] = _replace_emoji(b["text"])
        elif t == "quote":
            b["content"] = _replace_emoji(b["content"])
        elif t == "table":
            for row in b["rows"]:
                for i, c in enumerate(row):
                    row[i] = _replace_emoji(c)
        elif t == "code":
            b["content"] = _replace_emoji(b["content"])
    return blocks


# ============================================================
# 主流程
# ============================================================
def main():
    print(f"[1/4] 读取源文件 ...")
    md_script = SRC_SCRIPT.read_text(encoding="utf-8")
    md_qa = SRC_QA.read_text(encoding="utf-8")

    print(f"[2/4] 解析 Markdown ...")
    blocks_script = _clean_blocks(parse_markdown(md_script))
    blocks_qa = _clean_blocks(parse_markdown(md_qa))
    all_blocks = [
        ("第一部分 · 10 分钟答辩讲稿", blocks_script),
        ("第二部分 · 范式对比 + 超参调参 Q&A 弹药库", blocks_qa),
    ]

    print(f"[3/4] 生成 Word ...")
    gen_docx(all_blocks, OUT_DOCX)

    print(f"[4/4] 生成 PDF ...")
    gen_pdf(all_blocks, OUT_PDF, FONT_PATH)

    print(f"\n[DONE] 全部完成。")
    print(f"  - {OUT_DOCX}")
    print(f"  - {OUT_PDF}")


if __name__ == "__main__":
    main()
