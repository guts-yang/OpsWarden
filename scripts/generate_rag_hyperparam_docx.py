# -*- coding: utf-8 -*-
"""生成「RAG 超参数调优实验报告」Word 文档（图文并茂）。

数据来源：docs/rag_hyperparam_report.md (v2) 与 docs/rag_hyperparam_report_v3_joint.md (v3)。
依赖：python-docx, matplotlib。
输出：docs/RAG超参数调优实验报告.docx
"""
from pathlib import Path
import tempfile

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------- 品牌色 ----------------
ORANGE = RGBColor(0xFF, 0x67, 0x00)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
RED = RGBColor(0xC0, 0x39, 0x2B)
BLUE = RGBColor(0x1A, 0x73, 0xE8)
DARK = RGBColor(0x1A, 0x1A, 0x1F)
DIM = RGBColor(0x6B, 0x72, 0x80)

CN_FONT = "微软雅黑"
MONO_FONT = "Consolas"

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "RAG超参数调优实验报告.docx"

# matplotlib 中文
plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS"]
plt.rcParams["axes.unicode_minus"] = False

# ---------------- 数据 ----------------
TAUS = [0.35, 0.40, 0.45, 0.50, 0.55, 0.60, 0.65, 0.70, 0.75, 0.80, 0.85]
HIT_PARA = [100, 100, 100, 100, 100, 100, 98.5, 97.5, 95.5, 84.5, 66.5]
FPR = [96.9, 81.2, 62.5, 34.4, 15.6, 9.4, 3.1, 3.1, 0.0, 0.0, 0.0]


def make_charts(tmpdir: Path):
    c1 = tmpdir / "tau_curve.png"
    c2 = tmpdir / "default_vs_rec.png"

    # 图1：τ 权衡曲线
    fig, ax = plt.subplots(figsize=(8, 4.0), dpi=160)
    ax.plot(TAUS, HIT_PARA, "-o", color="#1e7e34", lw=2.5, ms=5, label="Hit@1_para（召回，越高越好）")
    ax.plot(TAUS, FPR, "-o", color="#c0392b", lw=2.5, ms=5, label="FPR（误命中，越低越好）")
    ax.fill_between(TAUS, HIT_PARA, color="#1e7e34", alpha=0.06)
    ax.fill_between(TAUS, FPR, color="#c0392b", alpha=0.06)
    # 高亮关键点
    ax.scatter([0.65], [98.5], s=130, facecolor="#1e7e34", edgecolor="white", zorder=5, linewidths=2)
    ax.scatter([0.40], [81.2], s=130, facecolor="#c0392b", edgecolor="white", zorder=5, linewidths=2)
    ax.annotate("推荐 τ=0.65\nHit 98.5% / FPR 3.1%", (0.65, 98.5), xytext=(0.66, 70),
                fontsize=9, color="#1e7e34", fontweight="bold",
                arrowprops=dict(arrowstyle="->", color="#1e7e34"))
    ax.annotate("默认 τ=0.40\nFPR 高达 81.2%", (0.40, 81.2), xytext=(0.42, 50),
                fontsize=9, color="#c0392b", fontweight="bold",
                arrowprops=dict(arrowstyle="->", color="#c0392b"))
    ax.axvspan(0.60, 0.70, color="#ff6700", alpha=0.07)
    ax.set_xlabel("相似度阈值 τ", fontsize=10)
    ax.set_ylabel("百分比 (%)", fontsize=10)
    ax.set_ylim(0, 108)
    ax.set_title("阈值 τ 权衡曲线（ε=0.02, L1-K=8, Top-K=3）", fontsize=11, fontweight="bold")
    ax.legend(loc="center left", fontsize=9, frameon=False)
    ax.grid(True, color="#e5e7eb", lw=0.8)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    fig.tight_layout()
    fig.savefig(c1, bbox_inches="tight")
    plt.close(fig)

    # 图2：默认 vs 推荐
    fig, ax = plt.subplots(figsize=(8, 3.8), dpi=160)
    labels = ["Hit@1_exact", "Hit@1_para", "FPR（越低越好）"]
    default = [100, 100, 81.2]
    rec = [100, 98.5, 3.1]
    x = range(len(labels))
    w = 0.36
    b1 = ax.bar([i - w / 2 for i in x], default, w, label="当前默认 (τ=0.40)", color="#c0392b")
    b2 = ax.bar([i + w / 2 for i in x], rec, w, label="推荐 (τ=0.65)", color="#1e7e34")
    for bars in (b1, b2):
        for b in bars:
            ax.text(b.get_x() + b.get_width() / 2, b.get_height() + 1.5,
                    f"{b.get_height():.1f}%", ha="center", va="bottom", fontsize=8.5, fontweight="bold")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, fontsize=10)
    ax.set_ylabel("百分比 (%)", fontsize=10)
    ax.set_ylim(0, 115)
    ax.set_title("当前默认 vs 推荐配置（τ 0.40 → 0.65）", fontsize=11, fontweight="bold")
    ax.legend(loc="upper right", fontsize=9, frameon=False)
    ax.grid(True, axis="y", color="#e5e7eb", lw=0.8)
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    fig.tight_layout()
    fig.savefig(c2, bbox_inches="tight")
    plt.close(fig)
    return c1, c2


# ---------------- docx 辅助 ----------------
def set_run(run, size=None, bold=None, color=None, font=CN_FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", size=10.5, bold=False, color=DARK, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{num}  ")
    set_run(r, size=15, bold=True, color=ORANGE)
    r2 = p.add_run(text)
    set_run(r2, size=15, bold=True, color=DARK)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=12, bold=True, color=DARK)
    return p


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def style_table(table, header_fill="FFF3E9"):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell in enumerate(table.rows[0].cells):
        shade_cell(cell, header_fill)
        for p in cell.paragraphs:
            for r in p.runs:
                set_run(r, size=9.5, bold=True, color=RGBColor(0xB5, 0x43, 0x0A))


def fill_table(table, rows, sizes=9.5):
    for i, row in enumerate(rows):
        cells = table.rows[i].cells
        for j, val in enumerate(row):
            cell = cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            txt, color, bold = val if isinstance(val, tuple) else (val, DARK, False)
            r = p.add_run(str(txt))
            set_run(r, size=sizes, bold=bold, color=color)


def add_mono_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(10)
        r = p.add_run(line if line else " ")
        r.font.name = MONO_FONT
        r._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(doc, lead, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(lead)
    set_run(r, size=10.5, bold=True, color=DARK)
    r2 = p.add_run(text)
    set_run(r2, size=10.5, bold=False, color=DARK)


# ---------------- 组装文档 ----------------
def build():
    doc = Document()
    doc.styles["Normal"].font.name = CN_FONT
    doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    doc.styles["Normal"].font.size = Pt(10.5)

    # 标题
    add_para(doc, "实验报告 · 答辩材料", size=9, bold=True, color=ORANGE, space_after=2)
    add_para(doc, "OpsWarden · RAG 超参数调优实验报告", size=20, bold=True, color=DARK, space_after=4)
    add_para(doc,
             "目标：在误命中率 FPR ≤ 5% 约束下，最大化知识库召回 Hit@1。"
             "评测集：exact + paraphrase + hard_confusion + negative；向量库：PostgreSQL / pgvector。",
             size=10, color=DIM, space_after=2)
    add_para(doc,
             "模型 BGE-small-zh-v1.5（512 维）  |  调参维度 ε · L1-K · τ · Top-K  |  网格组合 1584 (v2) / 216 (v3)",
             size=9.5, color=DIM, space_after=2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    tmp = Path(tempfile.mkdtemp())
    c1, c2 = make_charts(tmp)

    # 1 背景
    add_heading(doc, "1", "实验背景与目标")
    add_para(doc, "RAG 检索链由 4 个超参数控制，它们共同决定「问对了能不能查到、问偏了会不会乱答」：")
    t = doc.add_table(rows=5, cols=2)
    style_table(t)
    fill_table(t, [
        ["超参数", "作用"],
        [("ε  量化步长", BLUE, True), "向量吸附到网格，越大锚点合并越多，索引越轻"],
        [("L1-K  锚点路由 Top-K", BLUE, True), "粗筛阶段保留的候选锚点桶数量"],
        [("τ  相似度阈值", GREEN, True), "score ≥ τ 判命中，否则视为未命中并建工单"],
        [("Top-K  精排返回条数", BLUE, True), "最终送入 LLM 的知识上下文条数"],
    ])
    add_para(doc,
             "核心矛盾：阈值 τ 调低 → 召回高但误命中多（拿不相关知识硬答）；"
             "τ 调高 → 误命中少但可能漏答。调优就是找这个平衡点。",
             size=10, color=DIM, space_after=4)

    # 2 设置
    add_heading(doc, "2", "实验设置")
    t = doc.add_table(rows=5, cols=3)
    style_table(t)
    fill_table(t, [
        ["超参数", "候选范围（v2 网格）", "说明"],
        ["ANCHOR_QUANT_EPSILON (ε)", "0.01 – 0.10（6 档）", "量化网格步长"],
        ["RAG_ANCHOR_TOP_K (L1-K)", "1, 2, 3, 4, 8, 12, 16, 24", "锚点粗筛 Top-K"],
        ["RAG_SCORE_THRESHOLD (τ)", "0.35 – 0.85（步长 0.05）", "命中阈值"],
        ["RAG_TOP_K (Top-K)", "1, 3, 5", "精排返回条数"],
    ])
    add_para(doc,
             "评测集构造（v2，共 332 条）：原题 exact 100 + 口语改写 paraphrase 200 + 知识库外 negative 32。"
             "v3 进一步扩到 680 条（加 hard_confusion 同类干扰 80 + negative 100），知识库扩至 400 条以测锚点合并。",
             size=10, color=DIM)

    add_subheading(doc, "评测指标")
    t = doc.add_table(rows=5, cols=2)
    style_table(t)
    fill_table(t, [
        ["指标", "含义"],
        [("Hit@1_para", ORANGE, True), "主指标——口语改写问题的 Top-1 命中率，衡量「换种问法还能不能查到」"],
        [("FPR", RED, True), "误命中率——知识库外问题被误判为命中的比例，越低越好"],
        [("L1_recall", BLUE, True), "锚点召回——正确条目的锚点是否进入 L1 候选桶，验证粗筛不漏"],
        [("anchor_compression", GREEN, True), "锚点压缩比——量化后锚点数 / 条目数，越小索引越轻"],
    ])

    # 3 结果
    add_heading(doc, "3", "核心结果")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("一句话结论：")
    set_run(r, size=10.5, bold=True, color=ORANGE)
    r = p.add_run("最优组合 ε=0.02 · L1-K=8 · τ=0.65 · Top-K=3。其中 τ 从默认 0.4 提到 0.65 是关键——"
                  "几乎不损召回，却把误命中率从 81% 砍到 3%。")
    set_run(r, size=10.5, bold=False, color=DARK)

    doc.add_picture(str(c1), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc,
             "τ ≤ 0.5 时召回满分但 FPR 高得离谱（0.40 时 81.2%）；τ ≥ 0.75 后召回开始崩塌。"
             "τ=0.65 是甜点：Hit_para 仍有 98.5%，FPR 降到 3.1%，满足 ≤5% 约束。",
             size=10, space_after=8)

    doc.add_picture(str(c2), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "召回几乎不变，误命中率断崖式下降。", size=10, color=DIM, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_subheading(doc, "最优超参数组合")
    t = doc.add_table(rows=5, cols=4)
    style_table(t)
    fill_table(t, [
        ["参数", "当前默认（线上）", "v2 推荐", "v3 推荐（最新）"],
        ["ε（量化步长）", "0.02", "0.02", ("0.02", GREEN, True)],
        ["L1-K（锚点 Top-K）", "8", "8", ("8", GREEN, True)],
        ["τ（相似度阈值）", ("0.40", RED, True), "0.65", ("0.65", GREEN, True)],
        ["Top-K（精排条数）", "3", "3", ("3", GREEN, True)],
    ])

    add_subheading(doc, "推荐配置指标（v3，680 条评测 / 400 条扩库）")
    t = doc.add_table(rows=2, cols=4)
    style_table(t)
    fill_table(t, [
        ["Hit@1_exact", "Hit@1_para", "L1_recall_hard", "anchor_compression"],
        [("100%", GREEN, True), ("99.5%", ORANGE, True), ("96.2%", BLUE, True), ("0.25×", DARK, True)],
    ])
    add_para(doc,
             "关于 FPR 的两个数：v2（32 条负样本）测得 FPR≈3.1%，v3 扩到 100 条含 near-miss 后 FPR≈13%。"
             "数字升高不是模型变差，而是评测更严格；负样本量小、估计有方差，生产应持续用真实日志迭代。",
             size=10, color=DIM, space_after=6)

    add_subheading(doc, "Top 配置（v3 joint，按 Hit_para + L1_hard 选优）")
    t = doc.add_table(rows=5, cols=7)
    style_table(t)
    fill_table(t, [
        ["Rank", "ε", "L1-K", "Top-K", "Hit_para", "L1_hard", "FPR"],
        ["1", "0.02", "12", "3", "99.5%", "97.5%", "13.0%"],
        ["2", "0.01", "12", "3", "99.5%", "97.5%", "13.0%"],
        ["3", "0.02", "12", "1", "99.5%", "97.5%", "13.0%"],
        [("—", DIM, False), ("0.02", DARK, True), ("8", ORANGE, True), ("3", DARK, True), "99.5%", "96.2%", "13.0%"],
    ], sizes=9)
    add_para(doc,
             "统计最优 L1-K=12，但生产保守取 L1-K=8：召回仅差 1.3 个百分点，候选更少、延迟更低，"
             "且与现有默认一致、无需大改。",
             size=10, color=DIM)

    # 4 结论
    add_heading(doc, "4", "结论与生产建议")
    add_bullet(doc, "立即可落地：", "把 RAG_SCORE_THRESHOLD 从 0.40 调到 0.65，性价比最高。")
    add_bullet(doc, "其余三参数维持默认：", "ε=0.02、L1-K=8、Top-K=3 已是稳健解，无需改动。")
    add_bullet(doc, "显式写入 .env：", "当前 RAG_ANCHOR_TOP_K / ANCHOR_QUANT_EPSILON 未写进 .env，靠代码默认兜底，建议显式声明避免歧义。")
    add_bullet(doc, "持续迭代：", "负样本规模偏小，上线后应收集真实「未命中 / 误命中」日志，定期重跑调优脚本。")

    add_subheading(doc, "推荐 .env")
    add_mono_block(doc, [
        "# RAG 最优组合（v3 调优结论）",
        "ANCHOR_QUANT_EPSILON=0.02",
        "RAG_ANCHOR_TOP_K=8",
        "RAG_SCORE_THRESHOLD=0.65   # ← 由 0.40 上调，FPR 81% → 3%",
        "RAG_TOP_K=3",
    ])

    add_subheading(doc, "复现命令")
    add_mono_block(doc, [
        "python scripts/build_eval_dataset_v3.py",
        "python scripts/rag_hyperparam_eval.py --dataset v3 --joint3 --all",
        "python scripts/rag_hyperparam_verify_db.py --dataset v3 --joint3",
    ])

    doc.add_paragraph()
    add_para(doc,
             "数据来源：docs/rag_hyperparam_report.md (v2) · docs/rag_hyperparam_report_v3_joint.md (v3)。"
             "图表由 matplotlib 渲染。OpsWarden · AI 运维数字员工平台。",
             size=8.5, color=DIM)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    build()
